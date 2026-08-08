"""
Document parser service.

Extracts plain text + basic structure from supported formats:
  * PDF   -> pdfplumber (with PyPDF2 fallback); OCR fallback for scanned PDFs
  * DOCX  -> python-docx
  * TXT   -> direct read (utf-8, with gbk fallback)
  * MD    -> markdown text extracted, headings preserved
  * HTML  -> BeautifulSoup text extraction

Returns a ParsedDocument with pages (list of page texts), detected language
hint, title, and metadata. This is a pure function of the file; it never
calls any model.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional

from server.services.security import get_logger

logger = get_logger("parser")

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown", ".html", ".htm"}


@dataclass
class Page:
    index: int            # 1-based
    text: str
    char_count: int = 0


@dataclass
class ParsedDocument:
    path: Path
    file_type: str
    title: str = ""
    pages: List[Page] = field(default_factory=list)
    metadata: Dict[str, str] = field(default_factory=dict)
    language: str = "zh"
    needs_ocr: bool = False

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages if p.text.strip())

    @property
    def page_count(self) -> int:
        return len(self.pages)

    @property
    def char_count(self) -> int:
        return sum(p.char_count for p in self.pages)


# =====================================================================
# Public entry point
# =====================================================================
def parse_document(path: Path, ocr_service=None) -> ParsedDocument:
    """Parse a document by extension. Raises ValueError for unsupported types."""
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    logger.info("Parsing document: %s", path.name)
    if ext == ".pdf":
        doc = _parse_pdf(path, ocr_service)
    elif ext == ".docx":
        doc = _parse_docx(path)
    elif ext in (".txt",):
        doc = _parse_txt(path)
    elif ext in (".md", ".markdown"):
        doc = _parse_markdown(path)
    elif ext in (".html", ".htm"):
        doc = _parse_html(path)
    else:  # pragma: no cover - guarded above
        raise ValueError(f"Unsupported: {ext}")

    doc.language = _detect_language(doc.full_text)
    logger.info(
        "Parsed %s: %d pages, %d chars, needs_ocr=%s",
        path.name, doc.page_count, doc.char_count, doc.needs_ocr,
    )
    return doc


# =====================================================================
# Per-format parsers
# =====================================================================
def _parse_pdf(path: Path, ocr_service=None) -> ParsedDocument:
    pages: List[Page] = []
    title = path.stem
    extracted_text = True

    try:
        import pdfplumber

        with pdfplumber.open(str(path)) as pdf:
            title = (pdf.metadata or {}).get("Title", "") or path.stem
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if not text.strip():
                    extracted_text = False
                pages.append(Page(index=i, text=text.strip(), char_count=len(text)))
    except Exception as exc:  # noqa: BLE001
        logger.warning("pdfplumber failed (%s); trying PyPDF2 fallback", exc)
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(str(path))
            for i, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                if not text.strip():
                    extracted_text = False
                pages.append(Page(index=i, text=text.strip(), char_count=len(text)))
        except Exception as exc2:  # noqa: BLE001
            logger.error("PyPDF2 also failed: %s", exc2)
            pages = []

    # If we got essentially no text, this is likely a scanned PDF -> OCR.
    total_chars = sum(p.char_count for p in pages)
    if total_chars < 50 and ocr_service is not None:
        logger.info("Sparse text in PDF; running OCR via %s", type(ocr_service).__name__)
        try:
            ocr_pages = ocr_service.ocr_pdf(path)
            if ocr_pages:
                pages = [
                    Page(index=i + 1, text=t.strip(), char_count=len(t))
                    for i, t in enumerate(ocr_pages)
                ]
                extracted_text = True
        except Exception as exc:  # noqa: BLE001
            logger.error("OCR fallback failed: %s", exc)

    return ParsedDocument(
        path=path,
        file_type="pdf",
        title=title,
        pages=pages,
        needs_ocr=(not extracted_text),
    )


def _parse_docx(path: Path) -> ParsedDocument:
    from docx import Document as DocxDocument

    document = DocxDocument(str(path))
    paragraphs = [p.text for p in document.paragraphs]
    # Also extract tables.
    table_lines: List[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            table_lines.append(" | ".join(cells))
    text = "\n".join(paragraphs + table_lines)
    # DOCX has no real pagination; treat as one page but split by form-feed-ish.
    chunks = _split_long_text(text, target=3000)
    pages = [
        Page(index=i + 1, text=chunk, char_count=len(chunk))
        for i, chunk in enumerate(chunks)
    ]
    title = path.stem
    if document.core_properties.title:
        title = document.core_properties.title
    return ParsedDocument(
        path=path, file_type="docx", title=title, pages=pages or [Page(1, "")]
    )


def _parse_txt(path: Path) -> ParsedDocument:
    raw = _read_text_with_fallback(path)
    chunks = _split_long_text(raw, target=4000)
    pages = [
        Page(index=i + 1, text=c, char_count=len(c)) for i, c in enumerate(chunks)
    ]
    return ParsedDocument(
        path=path, file_type="txt", title=path.stem, pages=pages or [Page(1, "")]
    )


def _parse_markdown(path: Path) -> ParsedDocument:
    raw = _read_text_with_fallback(path)
    # Split on top-level headings to produce pseudo pages.
    parts = re.split(r"(?m)^(#{1,2}\s+.+)$", raw)
    pages: List[Page] = []
    current = ""
    idx = 1
    for part in parts:
        if not part:
            continue
        if len(current) + len(part) > 3500:
            pages.append(Page(index=idx, text=current.strip(), char_count=len(current)))
            idx += 1
            current = part
        else:
            current += part
    if current.strip():
        pages.append(Page(index=idx, text=current.strip(), char_count=len(current)))
    return ParsedDocument(
        path=path, file_type="markdown", title=path.stem, pages=pages or [Page(1, "")]
    )


def _parse_html(path: Path) -> ParsedDocument:
    from bs4 import BeautifulSoup

    raw = _read_text_with_fallback(path)
    soup = BeautifulSoup(raw, "html.parser")
    title = soup.title.string.strip() if soup.title and soup.title.string else path.stem
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    chunks = _split_long_text(text, target=4000)
    pages = [
        Page(index=i + 1, text=c.strip(), char_count=len(c)) for i, c in enumerate(chunks)
    ]
    return ParsedDocument(
        path=path, file_type="html", title=title, pages=pages or [Page(1, "")]
    )


# =====================================================================
# Helpers
# =====================================================================
def _read_text_with_fallback(path: Path) -> str:
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _split_long_text(text: str, target: int = 3000) -> List[str]:
    """Split text on paragraph boundaries into ~target-sized chunks."""
    if len(text) <= target:
        return [text] if text.strip() else []
    paragraphs = re.split(r"\n\s*\n", text)
    chunks: List[str] = []
    current = ""
    for para in paragraphs:
        if len(current) + len(para) > target and current:
            chunks.append(current.strip())
            current = para + "\n\n"
        else:
            current += para + "\n\n"
    if current.strip():
        chunks.append(current.strip())
    return [c for c in chunks if c]


def _detect_language(text: str) -> str:
    """Very lightweight CJK vs Latin detector (no model needed)."""
    if not text:
        return "en"
    cjk = sum(1 for ch in text[:2000] if "\u4e00" <= ch <= "\u9fff")
    return "zh" if cjk > 20 else "en"


def detect_section(page_text: str, max_len: int = 60) -> str:
    """Best-effort extraction of a heading/section label from page text."""
    if not page_text:
        return ""
    for line in page_text.splitlines():
        line = line.strip()
        if not line:
            continue
        # Markdown heading
        m = re.match(r"^#{1,4}\s+(.+)$", line)
        if m:
            return m.group(1)[:max_len]
        # 第X条 / 第X章 / 第X节 (Chinese legal numbering)
        if re.match(r"^第[一二三四五六七八九十百千0-9]+[条章节款项部分]", line):
            return line[:max_len]
        # Numbered headings like "1.2 概述" or "3. 付款方式"
        if re.match(r"^\d+(\.\d+)*[、.\s]\s*\S", line) and len(line) <= max_len:
            return line[:max_len]
    return ""
