"""
Report generator.

Produces Markdown / HTML / JSON / DOCX review reports from a DocumentAnalysis
result. Reports are saved under the user-isolated reports directory and
contain: document metadata, summary, risk list (with evidence, recommended
revision and legal clause reference), tender/technical-specific sections,
a compliance checklist, and an overall risk assessment.

No external document library is required for Markdown/HTML/JSON; DOCX uses
python-docx (already a dependency). Markdown is converted to HTML with the
`markdown` package for the demo UI.
"""
from __future__ import annotations

import json
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict

from server.config import Settings
from server.models.schemas import DocumentAnalysis, Report, RiskLevel
from server.services.security import get_logger, safe_user_dir

logger = get_logger("report")

_LEVEL_ORDER = {RiskLevel.HIGH: 0, RiskLevel.MEDIUM: 1, RiskLevel.LOW: 2}
_LEVEL_CN = {RiskLevel.HIGH: "高", RiskLevel.MEDIUM: "中", RiskLevel.LOW: "低"}
_LEVEL_COLOR = {
    RiskLevel.HIGH: "#e11d48",
    RiskLevel.MEDIUM: "#d97706",
    RiskLevel.LOW: "#16a34a",
}


def generate_report(
    analysis: DocumentAnalysis,
    settings: Settings,
    fmt: str = "markdown",
    user_id: str = "default",
) -> Report:
    fmt = fmt.lower()
    if fmt not in ("markdown", "html", "json", "docx"):
        raise ValueError(f"Unsupported report format: {fmt}")

    report_id = uuid.uuid4().hex[:12]
    title = f"DocGuard 审查报告 - {analysis.file_name}"
    reports_dir = safe_user_dir(settings, user_id, "reports")

    if fmt == "markdown":
        content = _render_markdown(analysis, title)
        ext = "md"
        file_path = reports_dir / f"docguard_report_{analysis.document_id}_{report_id}.{ext}"
        file_path.write_text(content, encoding="utf-8")
    elif fmt == "html":
        md = _render_markdown(analysis, title)
        content = _render_html(md, analysis, title)
        ext = "html"
        file_path = reports_dir / f"docguard_report_{analysis.document_id}_{report_id}.{ext}"
        file_path.write_text(content, encoding="utf-8")
    elif fmt == "docx":
        # DOCX is a binary format; the report body lives in the .docx file.
        file_path = _render_docx(analysis, title, reports_dir, report_id)
        content = ""
        ext = "docx"
    else:
        content = json.dumps(analysis.to_public_dict(), ensure_ascii=False, indent=2)
        ext = "json"
        file_path = reports_dir / f"docguard_report_{analysis.document_id}_{report_id}.{ext}"
        file_path.write_text(content, encoding="utf-8")

    logger.info("Report generated: %s", file_path)
    return Report(
        report_id=report_id,
        document_id=analysis.document_id,
        title=title,
        format=fmt,
        file_path=str(file_path),
        content=content,
    )


def _render_markdown(a: DocumentAnalysis, title: str) -> str:
    lines = []
    lines.append(f"# {title}")
    lines.append("")
    lines.append(f"> 生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"> DocGuard AI 本地文档智能审查 · 模型：{a.llm_model_name or '规则引擎（无LLM）'}")
    lines.append("")

    # Meta
    lines.append("## 一、文档基本信息")
    lines.append("")
    lines.append("| 项目 | 内容 |")
    lines.append("|------|------|")
    lines.append(f"| 文件名 | {a.file_name} |")
    lines.append(f"| 文档类型 | {a.summary.doc_type.value} |")
    lines.append(f"| 文件类型 | {a.file_type} |")
    lines.append(f"| 页数 | {a.page_count} |")
    lines.append(f"| 字符数 | {a.char_count} |")
    lines.append(f"| 文本块数 | {a.chunk_count} |")
    lines.append(f"| 综合风险等级 | **{_LEVEL_CN[a.overall_risk_level]}（{a.overall_risk_level.value}）** |")
    lines.append("")

    # Summary
    lines.append("## 二、文档摘要")
    lines.append("")
    if a.summary.summary_text:
        lines.append(a.summary.summary_text)
    elif a.summary.key_points:
        for kp in a.summary.key_points:
            lines.append(f"- {kp}")
    else:
        lines.append("（无摘要）")
    lines.append("")
    if a.summary.parties:
        lines.append("**相关方：** " + "、".join(a.summary.parties))
        lines.append("")

    # Risk overview
    counts = a.risk_count_by_level
    lines.append("## 三、风险概览")
    lines.append("")
    lines.append(f"- 高风险（High）：**{counts.get('High', 0)}** 项")
    lines.append(f"- 中风险（Medium）：**{counts.get('Medium', 0)}** 项")
    lines.append(f"- 低风险（Low）：**{counts.get('Low', 0)}** 项")
    lines.append("")

    # Risk list
    lines.append("## 四、风险明细")
    lines.append("")
    sorted_risks = sorted(a.risks, key=lambda r: _LEVEL_ORDER[r.risk_level])
    for r in sorted_risks:
        lines.append(f"### {r.id} [{_LEVEL_CN[r.risk_level]}] {r.issue}")
        lines.append("")
        lines.append(f"- **类别：** {r.category}")
        lines.append(f"- **位置：** {r.location}")
        lines.append(f"- **风险说明：** {r.explanation or '（无）'}")
        lines.append(f"- **修改建议：** {r.suggestion or '（无）'}")
        if r.clause:
            lines.append(f"- **法规依据：** {r.clause}")
        lines.append(f"- **原文证据：** ")
        lines.append(f"  > {r.evidence.replace(chr(10), ' ')}")
        lines.append("")

    # Compliance checklist
    if a.compliance_checklist:
        lines.append("## 五、合规检查清单")
        lines.append("")
        lines.append("| 检查项 | 状态 | 说明 |")
        lines.append("|--------|------|------|")
        for c in a.compliance_checklist:
            lines.append(f"| {c.item} | {c.status} | {c.note or '—'} |")
        lines.append("")

    # Tender section
    if a.requirements:
        lines.append("## 六、招标要求分析")
        lines.append("")
        if a.capability_match_score is not None:
            lines.append(f"**企业匹配度：{a.capability_match_score}%**")
            lines.append("")
        lines.append("| 编号 | 类别 | 要求 | 是否满足 |")
        lines.append("|------|------|------|----------|")
        for req in a.requirements:
            mark = "是" if req.matched else "待确认"
            lines.append(f"| {req.id} | {req.category} | {req.requirement[:80]} | {mark} |")
        lines.append("")
        if a.missing_capabilities:
            lines.append("**待确认/缺失能力：**")
            for m in a.missing_capabilities:
                lines.append(f"- {m}")
            lines.append("")

    # Technical section
    if a.chapter_checks:
        lines.append("## 七、技术方案章节检查")
        lines.append("")
        lines.append("| 章节 | 状态 | 说明 |")
        lines.append("|------|------|------|")
        for c in a.chapter_checks:
            mark = "已包含" if c.present else "缺失"
            lines.append(f"| {c.chapter} | {mark} | {c.note} |")
        lines.append("")
        if a.security_issues:
            lines.append("### 安全问题")
            lines.append("")
            for r in a.security_issues:
                lines.append(f"- **{r.issue}**（{r.location}）：{r.explanation}")
            lines.append("")
        if a.performance_risks:
            lines.append("### 性能风险")
            lines.append("")
            for r in a.performance_risks:
                lines.append(f"- **{r.issue}**（{r.location}）：{r.explanation}")
            lines.append("")

    # Footer
    lines.append("---")
    lines.append("")
    lines.append("本报告由 DocGuard AI 在本地生成，文档与模型均运行于本机，未上传任何云端。")
    if a.engine_notes:
        lines.append(f"引擎备注：{a.engine_notes}")
    return "\n".join(lines)


def _render_html(md: str, a: DocumentAnalysis, title: str) -> str:
    try:
        import markdown as md_lib

        body = md_lib.markdown(md, extensions=["tables", "fenced_code"])
    except Exception:  # noqa: BLE001
        body = "<pre>" + md.replace("<", "&lt;") + "</pre>"

    # Colorize the inline severity tags ([高]/[中]/[低]) produced by the markdown renderer.
    for lvl, cn in _LEVEL_CN.items():
        color = _LEVEL_COLOR[lvl]
        body = body.replace(
            f"[{cn}]",
            f'<span style="color:{color};font-weight:600;">[{cn}]</span>',
        )

    level = a.overall_risk_level.value
    color = _LEVEL_COLOR[a.overall_risk_level]
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8"/>
<title>{title}</title>
<style>
  body {{ font-family: "Microsoft YaHei","Segoe UI",sans-serif; max-width: 900px;
         margin: 40px auto; color: #1f2937; line-height: 1.7; padding: 0 20px; }}
  h1 {{ border-bottom: 3px solid #0f766e; padding-bottom: 10px; }}
  h2 {{ color: #0f766e; margin-top: 32px; border-left: 4px solid #0f766e; padding-left: 10px; }}
  table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
  th, td {{ border: 1px solid #d1d5db; padding: 8px 12px; text-align: left; }}
  th {{ background: #f3f4f6; }}
  blockquote {{ border-left: 4px solid #94a3b8; margin: 8px 0; padding: 4px 12px;
               color: #475569; background: #f8fafc; }}
  .badge {{ display:inline-block; padding:2px 10px; border-radius:12px; color:#fff;
           background:{color}; font-size:13px; }}
  .meta {{ color:#6b7280; font-size:13px; }}
</style>
</head>
<body>
<p class="meta">DocGuard AI 本地审查报告 · 综合风险等级 <span class="badge">{level}</span></p>
{body}
</body>
</html>"""


def _render_docx(a: DocumentAnalysis, title: str, reports_dir: Path, report_id: str) -> Path:
    """Generate a formatted .docx report (Word) — opens with tracked-changes-friendly layout."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    def rgb(hex_color: str) -> RGBColor:
        return RGBColor.from_string(hex_color.lstrip("#"))

    doc = Document()

    # Title
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
    meta.add_run(f"模型：{a.llm_model_name or '规则引擎（无LLM）'}").italic = True

    # ---- 基本信息 ----
    doc.add_heading("一、文档基本信息", level=1)
    meta_rows = [
        ("文件名", a.file_name),
        ("文档类型", a.summary.doc_type.value),
        ("文件类型", a.file_type),
        ("页数", str(a.page_count)),
        ("字符数", str(a.char_count)),
        ("文本块数", str(a.chunk_count)),
        ("综合风险等级", f"{_LEVEL_CN[a.overall_risk_level]}（{a.overall_risk_level.value}）"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    for k, v in meta_rows:
        cells = t.add_row().cells
        cells[0].text = k
        cells[1].text = v

    # ---- 摘要 ----
    doc.add_heading("二、文档摘要", level=1)
    if a.summary.summary_text:
        doc.add_paragraph(a.summary.summary_text)
    elif a.summary.key_points:
        for kp in a.summary.key_points:
            doc.add_paragraph(kp, style="List Bullet")
    else:
        doc.add_paragraph("（无摘要）")
    if a.summary.parties:
        p = doc.add_paragraph()
        p.add_run("相关方：").bold = True
        p.add_run("、".join(a.summary.parties))

    # ---- 风险概览 ----
    doc.add_heading("三、风险概览", level=1)
    counts = a.risk_count_by_level
    doc.add_paragraph(
        f"高风险：{counts.get('High', 0)} 项；中风险：{counts.get('Medium', 0)} 项；"
        f"低风险：{counts.get('Low', 0)} 项"
    )

    # ---- 风险明细 ----
    doc.add_heading("四、风险明细", level=1)
    sorted_risks = sorted(a.risks, key=lambda r: _LEVEL_ORDER[r.risk_level])
    for r in sorted_risks:
        rh = doc.add_heading(level=2)
        run = rh.add_run(f"{r.id} [{_LEVEL_CN[r.risk_level]}] {r.issue}")
        run.font.color.rgb = rgb(_LEVEL_COLOR[r.risk_level])
        for label, value in (
            ("类别", r.category),
            ("位置", r.location),
            ("风险说明", r.explanation or "（无）"),
            ("修改建议", r.suggestion or "（无）"),
        ):
            p = doc.add_paragraph()
            p.add_run(f"{label}：").bold = True
            p.add_run(value)
        if r.clause:
            p = doc.add_paragraph()
            p.add_run("法规依据：").bold = True
            p.add_run(r.clause)
        if r.evidence:
            p = doc.add_paragraph()
            p.add_run("原文证据：").bold = True
            p.add_run(r.evidence)

    # ---- 合规检查清单 ----
    if a.compliance_checklist:
        doc.add_heading("五、合规检查清单", level=1)
        ct = doc.add_table(rows=1, cols=3)
        ct.style = "Light Grid Accent 1"
        hdr = ct.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "检查项", "状态", "说明"
        for c in a.compliance_checklist:
            cells = ct.add_row().cells
            cells[0].text = c.item
            cells[1].text = c.status
            cells[2].text = c.note or "—"

    # ---- 招标要求 ----
    if a.requirements:
        doc.add_heading("六、招标要求分析", level=1)
        if a.capability_match_score is not None:
            p = doc.add_paragraph()
            p.add_run(f"企业匹配度：{a.capability_match_score}%").bold = True
        rt = doc.add_table(rows=1, cols=4)
        rt.style = "Light Grid Accent 1"
        hr = rt.rows[0].cells
        hr[0].text, hr[1].text, hr[2].text, hr[3].text = "编号", "类别", "要求", "是否满足"
        for req in a.requirements:
            cells = rt.add_row().cells
            cells[0].text = req.id
            cells[1].text = req.category
            cells[2].text = req.requirement[:80]
            cells[3].text = "是" if req.matched else "待确认"
        if a.missing_capabilities:
            p = doc.add_paragraph()
            p.add_run("待确认/缺失能力：").bold = True
            p.add_run("；".join(a.missing_capabilities))

    # ---- 技术方案章节 ----
    if a.chapter_checks:
        doc.add_heading("七、技术方案章节检查", level=1)
        ct2 = doc.add_table(rows=1, cols=3)
        ct2.style = "Light Grid Accent 1"
        hr2 = ct2.rows[0].cells
        hr2[0].text, hr2[1].text, hr2[2].text = "章节", "状态", "说明"
        for c in a.chapter_checks:
            cells = ct2.add_row().cells
            cells[0].text = c.chapter
            cells[1].text = "已包含" if c.present else "缺失"
            cells[2].text = c.note

    # ---- 页脚 ----
    doc.add_paragraph("")
    foot = doc.add_paragraph()
    foot.add_run(
        "本报告由 DocGuard AI 在本地生成，文档与模型均运行于本机，未上传任何云端。"
    ).italic = True
    if a.engine_notes:
        doc.add_paragraph(a.engine_notes)

    filename = f"docguard_report_{a.document_id}_{report_id}.docx"
    file_path = reports_dir / filename
    doc.save(str(file_path))
    return file_path
